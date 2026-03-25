from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


root = Path(r"c:\Users\xberi\Desktop\MISTICISMO-US")
book1_dir = root / "livros" / "livro_01"
source_path = book1_dir / "book1_the_shadow_beneath_the_craft_manuscript.txt"
output_path = book1_dir / "book1_the_shadow_beneath_the_craft_publish_ready.docx"

lines = source_path.read_text(encoding="utf-8").splitlines()

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)

title = ""
subtitle = ""
series = ""
author = ""

idx = 0
while idx < len(lines):
    line = lines[idx].strip()
    if line == "TITLE" and idx + 1 < len(lines):
        title = lines[idx + 1].strip()
        idx += 2
        continue
    if line == "SUBTITLE" and idx + 1 < len(lines):
        subtitle = lines[idx + 1].strip()
        idx += 2
        continue
    if line == "SERIES" and idx + 1 < len(lines):
        series = lines[idx + 1].strip()
        idx += 2
        continue
    if line == "AUTHOR" and idx + 1 < len(lines):
        author = lines[idx + 1].strip()
        idx += 2
        continue
    if line == "COPYRIGHT":
        break
    idx += 1

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(title)
r.bold = True
r.font.size = Pt(26)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(subtitle)
r.italic = True
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(series)
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(author)
r.font.size = Pt(12)

doc.add_page_break()

major_sections = {
    "PROLOGUE",
    "TABLE OF CONTENTS",
    "INTEGRATION COMPANION",
    "ACKNOWLEDGMENTS",
}

skip_labels = {"TITLE", "SUBTITLE", "SERIES", "AUTHOR"}

for raw in lines:
    line = raw.strip()
    if not line:
        doc.add_paragraph("")
        continue
    if line in skip_labels:
        continue
    if line in {"COPYRIGHT", "DEDICATION", "EPIGRAPH"}:
        doc.add_heading(line.title(), level=1)
        continue
    if line.startswith("Copyright ©"):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        continue
    if line == "All rights reserved.":
        doc.add_paragraph(line)
        continue
    if line in major_sections:
        doc.add_page_break()
        doc.add_heading(line.title(), level=1)
        continue
    if line.startswith("CHAPTER "):
        doc.add_page_break()
        doc.add_heading(line, level=1)
        continue
    if line.startswith("Integration Ritual —"):
        doc.add_heading(line, level=2)
        continue
    if line.startswith("Prompt "):
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.25)
        continue
    if line.startswith("- "):
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.25)
        continue
    if line[0].isdigit() and ")" in line[:4]:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.15)
        continue
    if line.endswith(":") and len(line) < 70:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.bold = True
        continue
    doc.add_paragraph(line)

doc.save(output_path)
print(str(output_path))
